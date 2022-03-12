# Importing necessary modules required
from playsound import playsound
import speech_recognition as sr
from googletrans import Translator
from gtts import gTTS
import os
import time
from pyfiglet import Figlet
from rich.progress import track
import docx2txt
from textblob import TextBlob
import nltk
from nltk import sent_tokenize
import docx


# main---


f = Figlet(font='slant')
print( f.renderText('--- TOVICE ---'))


print("Hello Everyone, This project is based on Voice Translation.")
txt = gTTS(text="Hello Everyone, This project is based on Voice Translation.", lang='en-us', slow=False)

txt.save("mytxt.mp3")
playsound('mytxt.mp3')
os.remove('mytxt.mp3')



while True:

        # Input from user
        print("\n#################### Available Options ####################")
        print()
        print("\n Type 'V' to operate in Voice Mode")
        print("\n Type 'M' to operate in Manual Mode")
        print("\n Type 'A' to see about Project Details")
        print("\n Type 'L' to see supported Languages for translation")
        print("\n Type 'T' to translate from a document file")
        print("\n Type 'R' to recognize text from given document")

        print("\n-------------------- Choose Option --------------------")
        print()
        usr = input("Choose one of the above option- ")
        option = ['V','v','M','m','about','About','Languages','languages','Language','language','T','t','R','r']

        while(usr not in option):
            print("Sorry the chose option is not supported, Choose again...")
            usr = input("Choose one of the above option- ")


        try:
           
              
            # A tuple containing all the language and
            # codes of the language will be detcted
            dic = ('afrikaans', 'af', 'albanian', 'sq', 
                   'amharic', 'am', 'arabic', 'ar',
                   'armenian', 'hy', 'azerbaijani', 'az', 
                   'basque', 'eu', 'belarusian', 'be',
                   'bengali', 'bn', 'bosnian', 'bs', 'bulgarian',
                   'bg', 'catalan', 'ca', 'cebuano',
                   'ceb', 'chichewa', 'ny', 'chinese (simplified)',
                   'zh-cn', 'chinese (traditional)',
                   'zh-tw', 'corsican', 'co', 'croatian', 'hr',
                   'czech', 'cs', 'danish', 'da', 'dutch',
                   'nl', 'english', 'en', 'esperanto', 'eo', 
                   'estonian', 'et', 'filipino', 'tl', 'finnish',
                   'fi', 'french', 'fr', 'frisian', 'fy', 'galician',
                   'gl', 'georgian', 'ka', 'german',
                   'de', 'greek', 'el', 'gujarati', 'gu',
                   'haitian creole', 'ht', 'hausa', 'ha',
                   'hawaiian', 'haw', 'hebrew', 'he', 'hindi',
                   'hi', 'hmong', 'hmn', 'hungarian',
                   'hu', 'icelandic', 'is', 'igbo', 'ig', 'indonesian', 
                   'id', 'irish', 'ga', 'italian',
                   'it', 'japanese', 'ja', 'javanese', 'jw',
                   'kannada', 'kn', 'kazakh', 'kk', 'khmer',
                   'km', 'korean', 'ko', 'kurdish (kurmanji)', 
                   'ku', 'kyrgyz', 'ky', 'lao', 'lo',
                   'latin', 'la', 'latvian', 'lv', 'lithuanian',
                   'lt', 'luxembourgish', 'lb',
                   'macedonian', 'mk', 'malagasy', 'mg', 'malay',
                   'ms', 'malayalam', 'ml', 'maltese',
                   'mt', 'maori', 'mi', 'marathi', 'mr', 'mongolian',
                   'mn', 'myanmar (burmese)', 'my',
                   'nepali', 'ne', 'norwegian', 'no', 'odia', 'or',
                   'pashto', 'ps', 'persian', 'fa',
                   'polish', 'pl', 'portuguese', 'pt', 'punjabi', 
                   'pa', 'romanian', 'ro', 'russian',
                   'ru', 'samoan', 'sm', 'scots gaelic', 'gd',
                   'serbian', 'sr', 'sesotho', 'st',
                   'shona', 'sn', 'sindhi', 'sd', 'sinhala', 'si',
                   'slovak', 'sk', 'slovenian', 'sl',
                   'somali', 'so', 'spanish', 'es', 'sundanese',
                   'su', 'swahili', 'sw', 'swedish',
                   'sv', 'tajik', 'tg', 'tamil', 'ta', 'telugu',
                   'te', 'thai', 'th', 'turkish',
                   'tr', 'ukrainian', 'uk', 'urdu', 'ur', 'uyghur',
                   'ug', 'uzbek',  'uz',
                   'vietnamese', 'vi', 'welsh', 'cy', 'xhosa', 'xh',
                   'yiddish', 'yi', 'yoruba',
                   'yo', 'zulu', 'zu')
              
              
            # Capture Voice
            # takes command through microphone
            def takecommand():  
                r = sr.Recognizer()
                with sr.Microphone() as source:
                    print("Say something.....")
                    r.adjust_for_ambient_noise(source, duration=1)
                    audio = r.listen(source)
              
                try:
                    print("Recognizing.....")
                    query = r.recognize_google(audio, language='en-in')
                    print(f"You said - {query}\n")

                except sr.RequestError as e:
                    print("Could not request results; {0}".format(e))

                except sr.UnknownValueError:
                    print("Error occurred")

                except Exception as e:
                    print("say that again please.....")

                return query
              

            if usr == "V" or usr == "v":
                print("\nYou chose Voice Mode...")
                time.sleep(1)
                query = takecommand()
                while (query == "None"):
                    query = takecommand()
                  
                  
                def destination_language():

                    print("\nChoose the language in which you want to convert : Ex. Hindi , English , etc.")
                    print()
                      
                    # Input destination language in
                    # which the user wants to translate
                    to_lang = takecommand()
                    while (to_lang == "None"):
                        to_lang = takecommand()
                    to_lang = to_lang.lower()
                    return to_lang
                  
                to_lang = destination_language()
                  
                # Mapping it with the code
                while (to_lang not in dic):
                    print("\nLanguage in which you are trying to convert is currently not available please input some other language")
                    print()
                    to_lang = destination_language()
                  
                to_lang = dic[dic.index(to_lang)+1]
                  
                  
                # invoking Translator
                translator = Translator()
                  
                  
                # Translating from src to dest
                text_to_translate = translator.translate(query, dest=to_lang)
                  
                text = text_to_translate.text
                  
                # Using Google-Text-to-Speech ie, gTTS() method
                # to speak the translated text into the
                # destination language which is stored in to_lang.
                # Also, we have given 3rd argument as False because
                # by default it speaks very slowly
                speak = gTTS(text=text, lang=to_lang, slow=False)
                  
                # Using save() method to save the translated
                # speech in capture_voice.mp3
                speak.save("captured_voice.mp3")
                  
                # Using OS module to run the translated voice.
                playsound('captured_voice.mp3')
                os.remove('captured_voice.mp3')
                  
                # Printing Output
                print(text)

                o = input("\nDo you want to save this in file- ")

                if o=="yes" or o=="y" or o=="Y" or o=="Yes":


                        try:

                                file_name=input("\nEnter the name of file- ")
                                f = file_name+".docx"

                                for step in track(range(100), description="Working..."):
                                        time.sleep(0.1)

                                doc = docx.Document()
                                doc.add_paragraph(text)
                                # save document
                                doc.save(f)

                                print("\nCreated Successfully.")

                        except FileExistsError:
                                print(f"\nSorry,file name {f} already exists")        


            elif usr == "M" or usr =="m" or usr=="manual":

                print("\nYou chose User Mode...")
                time.sleep(1)
                usrmd = input("\nType anything- ") 
                usrmd1 = input("\nType in which language do you want to translate it- ")

                for step in track(range(100), description=" Processing..."):
                    time.sleep(0.1)
                time.sleep(1)
                for step in track(range(100), description=" Hang on, Alomost done..."):
                    time.sleep(0.1)

                time.sleep(1)

                from_lang = usrmd.lower()
                des_lang = usrmd1.lower()

                
                des_lang = dic[dic.index(des_lang)+1]


                # invoking Translator
                translator = Translator()
                  
                # Translating from src to dest
                text_to_translate = translator.translate(from_lang, dest=des_lang)
                  
                text = text_to_translate.text
                  
                # Using Google-Text-to-Speech ie, gTTS() method
                # to speak the translated text into the
                # destination language which is stored in to_lang.
                # Also, we have given 3rd argument as False because
                # by default it speaks very slowly


                speak = gTTS(text=text, lang=des_lang, slow=False)
                  
                # Using save() method to save the translated
                # speech in capture_voice.mp3
                speak.save("captured_voice.mp3")
                  
                # Using OS module to run the translated voice.
                playsound('captured_voice.mp3')
                os.remove('captured_voice.mp3')
                  
                # Printing Output
                print(text)

                o = input("\nDo you want to save this in file- ")

                if o=="yes" or o=="y" or o=="Y" or o=="Yes":


                        try:

                                file_name=input("\nEnter the name of file- ")
                                f = file_name+".docx"

                                for step in track(range(100), description="Working..."):
                                        time.sleep(0.1)

                                doc = docx.Document()
                                doc.add_paragraph(text)
                                # save document
                                doc.save(f)

                                print("\nCreated Successfully.")

                        except FileExistsError:
                                print(f"\nSorry,file name {f} already exists")        



            elif usr == "A" or usr =="a" or usr =="about project" or usr =="About project" or usr =="About Project" or usr =="About" or usr =="about":

                for step in track(range(100), description="Loading Project Details..."):
                    time.sleep(0.1)

                print("\n-------------------- About Project --------------------")
                print()
                time.sleep(1)
                print("\nProject Name- Tovice")
                time.sleep(1)
                print("\nCreator - Rinku [CSE 3rd Year(2820902)]")
                time.sleep(1)
                print("\nLanguage Used - Python")
                time.sleep(1)
                print("\nProject Type - Command Line Based")
                time.sleep(1)
                print("\nProject Summary - ")
                print("\nTovice is a project with the help of which we can translate voices from one language to another.")
                print("\nThe project consists of two mode, User or Manual mode and Voice Mode for translation.")
                print("\nIn User Mode, user can directly operate the software by typing any sentence and then translate it.")
                print("\nIn Voice Mode, the software automatically take input from user, process it and the translate it in selected language choosen by the user.")


            elif usr == "Languages" or usr == "languages" or usr == "supported languages" or usr == "language" or usr=="Language":
                print()
                for step in track(range(100), description="Loading Supported Languages..."):
                    time.sleep(0.1)

                print("\n-------------------- List of Supported Languages --------------------")
                time.sleep(1)
                print()
                lang_support = ['afrikaans', 'albanian', 'amharic', 'arabic', 'armenian', 'azerbaijani', 'basque', 'belarusian', 'bengali', 'bosnian', 'bulgarian', 'catalan', 'cebuano', 'chichewa', 'chinese (simplified)', 'chinese (traditional)', 'corsican', 'croatian', 'czech', 'danish', 'dutch', 'english', 'esperanto', 'estonian', 'filipino', 'finnish', 'french', 'frisian', 'galician', 'georgian', 'german', 'greek', 'gujarati', 'haitian creole', 'hausa', 'hawaiian', 'hebrew', 'hindi', 'hmong', 'hungarian', 'icelandic', 'igbo', 'indonesian', 'irish', 'italian', 'japanese', 'javanese', 'kannada', 'kazakh', 'khmer', 'korean', 'kurdish (kurmanji)', 'kyrgyz', 'lao', 'latin', 'latvian', 'lithuanian', 'luxembourgish', 'macedonian', 'malagasy', 'malay', 'malayalam', 'maltese', 'maori', 'marathi', 'mongolian', 'myanmar (burmese)', 'nepali', 'norwegian', 'odia', 'pashto', 'persian', 'polish', 'portuguese', 'punjabi', 'romanian', 'russian', 'samoan', 'scots gaelic', 'serbian', 'sesotho', 'shona', 'sindhi', 'sinhala', 'slovak', 'slovenian', 'somali', 'spanish', 'sundanese', 'swahili', 'swedish', 'tajik', 'tamil', 'telugu', 'thai', 'turkish', 'ukrainian', 'urdu', 'uyghur', 'uzbek', 'vietnamese', 'welsh', 'xhosa', 'yiddish', 'yoruba', 'zulu']
                a=0
                for i in lang_support:
                    a+=1
                    print(f"({a})",i.capitalize())



            elif usr == "t" or usr == "T":

                try:

                        file_name = input("\nEnter the name of the file-")

                        new_file = file_name+".docx"
         

                        text1 = docx2txt.process(new_file)
                        
                        des_lang = input("\nChoose language in which you want to convert it- ")
                        des_lang = des_lang.lower()
                        des_lang = dic[dic.index(des_lang)+1]


                        # invoking Translator
                        translator = Translator()
                          
                        # Translating from src to dest
                        text_to_translate = translator.translate(text1, dest=des_lang)
                          
                        text = text_to_translate.text

                        speak = gTTS(text=text, lang=des_lang, slow=False)
                          
                        # Using save() method to save the translated
                        # speech in capture_voice.mp3
                        speak.save("captured_voice1.mp3")
                          
                        # Using OS module to run the translated voice.
                        playsound('captured_voice1.mp3')
                        os.remove('captured_voice1.mp3')
                          
                        # Printing Output
                        print(text)

                        o = input("\nDo you want to save this in file- ")

                        if o=="yes" or o=="y" or o=="Y" or o=="Yes":

                                try:
                                

                                        file_name=input("\nEnter the name of file- ")
                                        f = file_name+".docx"

                                        for step in track(range(100), description="Working..."):
                                                time.sleep(0.1)

                                        doc = docx.Document()
                                        doc.add_paragraph(text)
                                        # save document
                                        doc.save(f)

                                        txt = gTTS(text="Created Successfully", lang='en-us', slow=False)

                                        txt.save("mytxt1.mp3")
                                        playsound('mytxt1.mp3')
                                        os.remove('mytxt1.mp3')

                                        print("\nCreated Successfully.")


                                except FileExistsError:
                                        print(f"\nSorry,file name {f} already exists")        

                except FileNotFoundError:
                        print("\nSorry file not found:(")

            elif usr == "R" or usr == "r":

                try:

                        file_name = input("\nEnter the name of the file-")

                        new_file = file_name+".docx"

                        my_text = docx2txt.process(new_file)
                        text = my_text
                        n = sent_tokenize(text)
                        n1 = str(n[0])
                        translator = Translator()
                        a = translator.detect(n1)
                        l = a.lang
                        if l in dic:
                                for step in track(range(100), description="Detecting..."):
                                        time.sleep(0.1)
                                print(dic[dic.index(l)-1])
                
                except FileNotFoundError:
                        print("\nSorry file not found:(")

            else:
                print("\nSorry the command is not supported...")        

            

        except KeyboardInterrupt:
                print("\nCtrl + C Key is pressed...")
                time.sleep(1)
                print("\nExiting...")  


        except UnboundLocalError:

            print("\nSorry, didn't get you. Speak again...")


        again = input("\nDo you want to continue?(Y/N)")
        if again == "N" or again == "n" or again == "no":
                break
